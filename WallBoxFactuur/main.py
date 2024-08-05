#### Import

from docx import Document
from docx.shared import Inches
from wallbox import Wallbox, Statuses

import datetime
from datetime import datetime
import pandas as pd
import math
import matplotlib.pyplot as plt
import PySimpleGUI as sg

###### variabelen ######
Year_Report  = datetime.now().year
Month_Report = datetime.now().month-1
chargerId    = 26049
Tariff       = 0.35
w = Wallbox("koob@koobnobbe.com", "SuperHero@01")

def get_wallbox_data(year,month):
    ####### Load data #########
    startdate_str = str(year) + '-' + str(month) + '-1'
    if month ==0:
        startdate_str = str(year-1) + '-' + str(12) + '-1'
        enddate_str = str(year) + '-' + str(month + 1 ) + '-1'
    else:
        startdate_str = str(year) + '-' + str(month) + '-1'
        enddate_str = str(year) + '-' + str(month + 1) + '-1'

    startDate = datetime.strptime(startdate_str, '%Y-%m-%d')
    #enddate_str = '2023-1-1'
    endDate = datetime.strptime(enddate_str, '%Y-%m-%d')
    w.authenticate()

    df = pd.json_normalize(w.getSessionList(chargerId, startDate, endDate))
    df2 = pd.json_normalize(df['data']).T
    df3 = pd.json_normalize(df2[0]).copy(deep=True)
    df3.reset_index(inplace=True)
    df4 = df3[['attributes.start', 'attributes.end', 'attributes.energy', 'attributes.energy_unit'
               ]].copy(deep=True)
    ####### Converst datatypes ###############
    df4['startdate'] = df4['attributes.start'].astype(int).apply(
        lambda x: datetime.fromtimestamp(x).strftime('%Y-%m-%d %H:%M:%S'))
    df4['startdate'] = pd.to_datetime(df4['startdate'], format='%Y-%m-%d %H:%M:%S')

    df4['enddate'] = df4['attributes.end'].astype(int).apply(
        lambda x: datetime.fromtimestamp(x).strftime('%Y-%m-%d %H:%M:%S'))
    df4['enddate'] = pd.to_datetime(df4['enddate'], format='%Y-%m-%d %H:%M:%S')

    df4.drop(['attributes.start', 'attributes.end'], axis=1, inplace=True)
    df4 = df4.rename(columns={'attributes.energy': 'energy', 'attributes.energy_unit': 'unit'})
    df4 = df4[['startdate', 'enddate', 'energy', 'unit']]

    ###### Caclulations #######
    df4['hour'] = pd.to_datetime(df4['startdate']).dt.hour
    df4['Datum'] = pd.to_datetime(df4['startdate']).dt.date
    df4['Tijd'] = pd.to_datetime(df4['startdate']).dt.time
    df4['Maand'] = pd.to_datetime(df4['startdate']).dt.month
    df4['Omschrijving'] = 'Afgenomen energie (kwh)'
    df4['Tarief'] = Tariff
    df4['Hoeveelheid (kwh)'] = df4['energy']
    df4['Hoeveelheid (Euro)'] = (df4['energy'] * Tariff).round(2)
    df4.drop(['startdate', 'hour', 'energy', 'enddate', 'unit'], axis=1, inplace=True)

    #####
    df_WallBoxEnergie = df4.copy(deep=True)
    ######
    return df_WallBoxEnergie
def add_staticcosts(df_wb):
    row_static = pd.DataFrame({'Datum': [math.nan],
                               'Tijd': [math.nan],
                               'Maand': [Month_Report],
                               'Omschrijving': ['Vaste verbruikskosten Kantoor'],
                               'Tarief': [15],
                               'Hoeveelheid (kwh)': [1],
                               'Hoeveelheid (Euro)': [15]})
    df_ttl = pd.concat([df_wb, row_static], axis=0)
    return df_ttl
def get_wallbox_totals(df):
    return df.groupby(['Maand','Omschrijving','Tarief']).sum(['Hoeveelheid (kwh)','Hoeveelheid (EURO)']).reset_index()
def get_wallbox_generaltotal(df):
    return df.groupby(['Maand']).sum(['Hoeveelheid (kwh)','Hoeveelheid (EURO)']).reset_index()

def get_col_length(type='text', column=0, df=[], text='', textsize=11):
    length= 0
    if type == 'df':
        length = len(df[column])
    if type == 'text':
        length = len(text)
    return int(round(length*textsize/3,0))

def df_to_table(df_detail):
    from matplotlib.backends.backend_pdf import PdfPages
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.axis('tight')
    ax.axis('off')
    the_table = ax.table( cellText=df_detail.values
                         , colLabels=df_detail.columns
                         , loc='center')
    return the_table

def find_replace(paragraph_keyword, draft_keyword, paragraph):
    #print(paragraph.text)
    if paragraph_keyword in paragraph.text:
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)

def create_invoice(df, filename):
    df_all = add_staticcosts(df)
    df_ttl = get_wallbox_totals(df_all)
    df_gt = get_wallbox_generaltotal(df_all)
    ########################################
    document = Document('template.docx')
    style= 'Table Grid'
    document.add_paragraph('Overzicht van de kosten :')

    [document.add_paragraph('') for _ in range(1)]

    tableCols = 4
    table_ttl = document.add_table(rows=1, cols=tableCols)
    table_ttl.style = style
    hdr_cells = table_ttl.rows[0].cells
    hdr_cells[0].text = 'Omschrijving'
    hdr_cells[1].text = 'Tarief'
    hdr_cells[2].text = 'Hoeveelheid (kwh)'
    hdr_cells[3].text = 'Prijs'
    for i in range(tableCols):
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    table_ttl.cell(0,0).width = Inches(4)
    for i in range(len(df_ttl.index)):
        row_cells = table_ttl.add_row().cells
        row_cells[0].text = df_ttl.iloc[i, 1]  ##Omschrijving
        row_cells[1].text = str(df_ttl.iloc[i, 2])  ##Tarief
        row_cells[2].text = '%.1f' % df_ttl.iloc[i, 3]
        row_cells[3].text = '€ %.2f' % df_ttl.iloc[i, 4]

    row_cells = table_ttl.add_row().cells
    row_cells[0].text ='Totaal'
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = '€ %.2f' % df_gt.iloc[0, 3]

    [document.add_paragraph('') for _ in range(2)]
    print(df_gt)
    document.add_paragraph('Gelieve € %.2f te betalen binnen 14 dagen na dagtekening op rekening: 39.66.34.397 t.n.v. K.Nobbe o.v.v. faktuurnummer' % df_gt.iloc[0, 3])
    document.add_page_break()

    document.add_paragraph('Details:')
    table = document.add_table(rows=1, cols=5 )
    table.style = style
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Datum'
    hdr_cells[1].text = 'Omschrijving'
    hdr_cells[2].text = 'Tarief'
    hdr_cells[3].text = 'Hoeveelheid (kwh)'
    hdr_cells[4].text = 'Prijs'
    for i in range(5):
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    for i in range(len(df.index)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(df.iloc[i,0]) ## Datum
        row_cells[1].text = df.iloc[i,3]      ##Omschrijving
        row_cells[2].text = str(df.iloc[i,4]) ##Tarief
        row_cells[3].text = '%.1f' % df.iloc[i,5]
        row_cells[4].text = '€ %.2f' %  df.iloc[i,6]
    ## Retrns 2x
    [document.add_paragraph('') for _ in range(2)]

    for paragraph in document.paragraphs:
        find_replace("%%YEAR%%", str(Year_Report), paragraph)
        find_replace("%%MONTH%%", str(Month_Report), paragraph)
        find_replace("%%DATE%%", datetime.now().strftime('%d-%m-%Y') , paragraph)
    document.save(filename)
    return 0


def create_pdf(filename):
    from docx2pdf import convert
    convert(filename)
##### Main Program ######

def Show_Parameter_Dialog(year_report, month_report, tariff):
    title = 'Set Parameters'
    msg = 'Give the month of reporting:'
    year_list = ['2020','2021','2022','2023','2024','2025']
    month_list = ['1','2','3','4','5','6','7','8','9','10','11','12']

    sg.set_options(border_width=0, margins=(0, 0), element_padding=(5, 3))
    sg.theme('Reddit')  # Add a touch of color
    font = ("Arial", 24)

    layout = [  [sg.Text(msg, font=font)],
                [sg.Text(' Year of reporting: ', font=font ), sg.Spin(values = year_list,initial_value=year_report, key='-YEAR-IN-', font=font, size=(10,1))],
                [sg.Text(' Month of reporting: ', font=font), sg.Spin(values= month_list, initial_value=month_report, key='-MONTH-IN-', font=font, size=(8,1))],
                [sg.Text(' Tariff : ', font=font), sg.InputText(default_text=tariff,size=(5,1), key='-TARIFF-IN-', font=font) ],
                [sg.Button('Process', font=font), sg.Exit( font=font)],
             ]
    window = sg.Window(title, layout,  finalize=True)
    action = ''
    while True:  # The Event Loop
        event, values = window.read()
        if event == 'Process':
            year_report = int(values['-YEAR-IN-'])
            month_report = int(values['-MONTH-IN-'])
            tariff = float(values['-TARIFF-IN-'])
            action = 'process'
            break

        if event == sg.WIN_CLOSED or event == 'Exit':
            action = 'exit'
            break

    window.close()
    return year_report, month_report, tariff, action


if __name__ == '__main__':
    print('Start...')
 #   Year_Report,Month_Report,Tariff, Action = Show_Parameter_Dialog(Year_Report,Month_Report,Tariff)
    print(Year_Report, Month_Report, Tariff)
 #   if Action == 'process':
    df_detail = get_wallbox_data(Year_Report,Month_Report)
    filename = 'Factuur_Wallbox'+str(Year_Report)+str(Month_Report)+'.docx'
    create_invoice(df_detail,filename)
    create_pdf(filename)
    #if Action == 'exit':
    #    print('No Futher Processing!')

    print('...Done!')